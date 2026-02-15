"""
Report Generator Service for ComplianceFlow.

Generates compliance reports with AI-interpreted content in multiple formats
(Markdown, PDF, DOCX, XLSX). Uses OllamaService for content generation
with graceful fallback when Ollama is unavailable.

Architecture:
- Pure functions for prompt building and format conversion (no side effects)
- Async orchestrator for AI content generation (single side effect: Ollama call)
- All functions are fully typed with docstrings
"""

import json
import logging
from typing import Any, Dict, List

from app.services.ollama import OllamaService, CompletionRequest

logger = logging.getLogger(__name__)


# ============================================================================
# Constants
# ============================================================================

MAX_INPUT_CHARS: int = 12000  # Truncate input data to fit context window
DEFAULT_MODEL: str = "llama3.2:3b"


# ============================================================================
# Pure Functions — Content Generation
# ============================================================================


def _serialize_input_data(input_data: Dict[str, Any], max_chars: int = MAX_INPUT_CHARS) -> str:
    """Serialize input data to a truncated string representation.

    Pure function. Converts dict to formatted JSON, truncating if needed.

    Args:
        input_data: The upstream workflow data to serialize.
        max_chars: Maximum characters to include.

    Returns:
        A string representation of the data, truncated with notice if needed.
    """
    try:
        raw: str = json.dumps(input_data, indent=2, default=str)
    except (TypeError, ValueError):
        raw = str(input_data)

    if len(raw) > max_chars:
        return raw[:max_chars] + "\n\n... [Data truncated for context window]"
    return raw


def _build_fallback_markdown(
    input_data: Dict[str, Any],
    frameworks: List[str],
    report_title: str,
) -> str:
    """Build a structured markdown report directly from input data (no AI).

    Pure function. Used as fallback when Ollama is unavailable.

    Args:
        input_data: The upstream workflow data.
        frameworks: List of compliance framework names.
        report_title: Title for the report.

    Returns:
        A markdown string with structured sections.
    """
    title: str = report_title or "Compliance Report"
    fw_list: str = ", ".join(frameworks) if frameworks else "General"

    sections: List[str] = [
        f"# {title}",
        "",
        f"**Frameworks**: {fw_list}",
        "",
        "## Executive Summary",
        "",
        "This report was generated without AI interpretation (Ollama unavailable).",
        "The raw data from the compliance workflow is presented below.",
        "",
        "## Data Summary",
        "",
    ]

    if input_data:
        for key, value in input_data.items():
            # Skip internal keys
            if key.startswith("_"):
                continue
            display_value: str = str(value)[:500]
            sections.append(f"### {key}")
            sections.append("")
            sections.append(f"```\n{display_value}\n```")
            sections.append("")
    else:
        sections.append("No input data was provided to this compliance node.")
        sections.append("")

    return "\n".join(sections)


def build_report_prompt(
    input_data: Dict[str, Any],
    frameworks: List[str],
    report_title: str,
    include_evidence: bool,
) -> str:
    """Construct the LLM prompt from input data.

    Pure function. No side effects.

    Args:
        input_data: The upstream workflow data to interpret.
        frameworks: List of compliance framework names (e.g., ["SOC2", "GDPR"]).
        report_title: Title for the compliance report.
        include_evidence: Whether to include evidence references in the report.

    Returns:
        A structured prompt string for the LLM.
    """
    title: str = report_title or "Compliance Report"
    fw_list: str = ", ".join(frameworks) if frameworks else "General compliance"
    serialized_data: str = _serialize_input_data(input_data)

    prompt_parts: List[str] = [
        f"You are a compliance analyst. Generate a professional compliance report titled '{title}'.",
        f"",
        f"Frameworks to assess against: {fw_list}",
        f"",
        f"Analyze the following data from a compliance workflow and produce a well-structured "
        f"Markdown report with these sections:",
        f"1. Executive Summary — key findings in 2-3 paragraphs",
        f"2. Compliance Assessment — detailed analysis per framework",
        f"3. Findings — specific compliance gaps or confirmations",
        f"4. Recommendations — actionable items to improve compliance posture",
    ]

    if include_evidence:
        prompt_parts.append(
            "5. Evidence References — include specific data points from the input as evidence"
        )

    prompt_parts.extend([
        "",
        "Use proper Markdown formatting with headers, bullet points, and tables where appropriate.",
        "",
        "## Input Data",
        "",
        serialized_data,
    ])

    return "\n".join(prompt_parts)


# ============================================================================
# Async Functions — AI Content Generation
# ============================================================================


async def generate_report_content(
    input_data: Dict[str, Any],
    frameworks: List[str],
    report_title: str,
    include_evidence: bool,
    model: str = DEFAULT_MODEL,
) -> str:
    """Generate AI-interpreted compliance report content as Markdown.

    Calls OllamaService.generate() with the constructed prompt.
    Gracefully falls back to structured markdown if Ollama is unavailable.

    Args:
        input_data: The upstream workflow data to interpret.
        frameworks: List of compliance framework names.
        report_title: Title for the compliance report.
        include_evidence: Whether to include evidence references.
        model: Ollama model name to use.

    Returns:
        A Markdown string containing the report content.
    """
    prompt: str = build_report_prompt(
        input_data=input_data,
        frameworks=frameworks,
        report_title=report_title,
        include_evidence=include_evidence,
    )

    try:
        ollama = OllamaService()
        request = CompletionRequest(
            model=model,
            prompt=prompt,
            temperature=0.3,  # Low temperature for factual reporting
        )
        response = await ollama.generate(request)

        content: str = response.response.strip()
        if content:
            return content

        # Empty response — fall back
        logger.warning("Ollama returned empty response, using fallback")

    except Exception as e:
        logger.warning(f"Ollama unavailable ({type(e).__name__}: {e}), using fallback")

    # Fallback: generate structured markdown without AI
    return _build_fallback_markdown(
        input_data=input_data,
        frameworks=frameworks,
        report_title=report_title,
    )


# ============================================================================
# Orchestrator — Report Generation Pipeline
# ============================================================================

# Format → (generator, mime_type, extension)
FORMAT_REGISTRY: Dict[str, tuple] = {
    "md": ("text/markdown", ".md"),
    "pdf": ("application/pdf", ".pdf"),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"),
    "xlsx": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", ".xlsx"),
}

SUPPORTED_FORMATS: List[str] = list(FORMAT_REGISTRY.keys())


def _format_to_generator(report_format: str) -> str:
    """Normalize format string and validate.

    Pure function.

    Args:
        report_format: Requested format string (case-insensitive).

    Returns:
        Normalized format key from FORMAT_REGISTRY, defaulting to 'md'.
    """
    normalized: str = report_format.lower().strip()
    if normalized not in FORMAT_REGISTRY:
        logger.warning(f"Unsupported format '{report_format}', falling back to markdown")
        return "md"
    return normalized


async def generate_report(
    input_data: Dict[str, Any],
    report_format: str,
    frameworks: List[str],
    report_title: str,
    include_evidence: bool,
    model: str = DEFAULT_MODEL,
    report_prompt: str = "",
) -> Dict[str, Any]:
    """Orchestrate full report generation pipeline.

    Calls generate_report_content for AI content, then the appropriate format generator.
    Returns a dict with base64-encoded file content, filename, and MIME type.

    Args:
        input_data: The upstream workflow data to interpret.
        report_format: Desired output format (md, pdf, docx, xlsx).
        frameworks: List of compliance framework names.
        report_title: Title for the compliance report.
        include_evidence: Whether to include evidence references.
        model: Ollama model name to use.
        report_prompt: Optional user-provided custom prompt to guide AI interpretation.

    Returns:
        Dict with keys: file_content (base64), filename, mime_type, format, success.
    """
    import base64

    fmt: str = _format_to_generator(report_format)
    mime_type, extension = FORMAT_REGISTRY[fmt]
    title: str = report_title or "Compliance Report"

    # If user provided a custom prompt, prepend it to input_data
    effective_data: Dict[str, Any] = dict(input_data)
    if report_prompt:
        effective_data["_user_report_prompt"] = report_prompt

    # Step 1: Generate AI content
    content: str = await generate_report_content(
        input_data=effective_data,
        frameworks=frameworks,
        report_title=title,
        include_evidence=include_evidence,
        model=model,
    )

    # Step 2: Convert to requested format
    if fmt == "md":
        file_bytes: bytes = generate_markdown(content=content, title=title)
    elif fmt == "pdf":
        file_bytes = generate_pdf(content=content, title=title)
    elif fmt == "docx":
        file_bytes = generate_docx(content=content, title=title)
    elif fmt == "xlsx":
        file_bytes = generate_xlsx(content=content, title=title, input_data=input_data)
    else:
        file_bytes = generate_markdown(content=content, title=title)

    # Step 3: Encode and return
    safe_title: str = title.replace(" ", "_").replace("/", "-")[:50]
    filename: str = f"{safe_title}{extension}"

    return {
        "file_content": base64.b64encode(file_bytes).decode("utf-8"),
        "filename": filename,
        "mime_type": mime_type,
        "format": fmt,
        "success": True,
    }


# ============================================================================
# Pure Functions — Format Generators
# ============================================================================


def generate_markdown(content: str, title: str) -> bytes:
    """Wrap content with a title header and return as UTF-8 bytes.

    Pure function. No side effects.

    Args:
        content: Markdown content body.
        title: Report title (rendered as H1 header).

    Returns:
        UTF-8 encoded bytes of the complete markdown document.
    """
    title_str: str = title or "Compliance Report"
    full_doc: str = f"# {title_str}\n\n{content}"
    return full_doc.encode("utf-8")


def generate_pdf(content: str, title: str) -> bytes:
    """Convert markdown content to a styled PDF using reportlab.

    Uses Nomu brand fonts (Barlow for headings, Open Sans for body) and
    brand colors (primary #4004DA, accent #FF6C1D).

    Args:
        content: Markdown content body.
        title: Report title.

    Returns:
        PDF file as bytes.
    """
    import io as _io
    import os as _os
    import re as _re
    from pathlib import Path as _Path
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        HRFlowable,
        PageBreak,
        KeepTogether,
    )
    from reportlab.graphics.shapes import Drawing, Line

    # ── Nomu Brand Colors ────────────────────────────────────────────
    NOMU_PRIMARY = HexColor("#4004DA")       # Purple-blue
    NOMU_PRIMARY_LIGHT = HexColor("#EDE9F8") # Very light purple tint
    NOMU_BLACK = HexColor("#1A1A2E")         # Near-black
    NOMU_WHITE = HexColor("#FFFFFF")         # White
    NOMU_SURFACE = HexColor("#F8F8FA")       # Light surface
    NOMU_MUTED = HexColor("#9CA3AF")         # Muted text
    NOMU_DARK = HexColor("#374151")          # Dark gray for headings
    NOMU_BORDER = HexColor("#E5E7EB")        # Border gray
    NOMU_TABLE_ALT = HexColor("#FAFAFA")     # Alternating row

    # ── Font Registration ────────────────────────────────────────────
    fonts_dir = _Path(__file__).parent.parent / "assets" / "fonts"

    # Heading font: Barlow
    HEADING_FONT = "Helvetica"
    HEADING_FONT_BOLD = "Helvetica-Bold"
    # Body font: Open Sans
    BODY_FONT = "Helvetica"
    BODY_FONT_BOLD = "Helvetica-Bold"
    BODY_FONT_ITALIC = "Helvetica-Oblique"
    BODY_FONT_BOLDITALIC = "Helvetica-BoldOblique"

    try:
        if (fonts_dir / "Barlow-Regular.ttf").exists():
            pdfmetrics.registerFont(TTFont("Barlow", str(fonts_dir / "Barlow-Regular.ttf")))
            pdfmetrics.registerFont(TTFont("Barlow-SemiBold", str(fonts_dir / "Barlow-SemiBold.ttf")))
            pdfmetrics.registerFont(TTFont("Barlow-Bold", str(fonts_dir / "Barlow-Bold.ttf")))
            HEADING_FONT = "Barlow"
            HEADING_FONT_BOLD = "Barlow-Bold"

        if (fonts_dir / "OpenSans-Regular.ttf").exists():
            pdfmetrics.registerFont(TTFont("OpenSans", str(fonts_dir / "OpenSans-Regular.ttf")))
            pdfmetrics.registerFont(TTFont("OpenSans-Bold", str(fonts_dir / "OpenSans-Bold.ttf")))
            pdfmetrics.registerFont(TTFont("OpenSans-Italic", str(fonts_dir / "OpenSans-Italic.ttf")))
            pdfmetrics.registerFont(TTFont("OpenSans-BoldItalic", str(fonts_dir / "OpenSans-BoldItalic.ttf")))
            pdfmetrics.registerFontFamily(
                "OpenSans",
                normal="OpenSans",
                bold="OpenSans-Bold",
                italic="OpenSans-Italic",
                boldItalic="OpenSans-BoldItalic",
            )
            BODY_FONT = "OpenSans"
            BODY_FONT_BOLD = "OpenSans-Bold"
            BODY_FONT_ITALIC = "OpenSans-Italic"
            BODY_FONT_BOLDITALIC = "OpenSans-BoldItalic"
    except Exception:
        pass  # Fall back to Helvetica

    # ── Helpers ───────────────────────────────────────────────────────
    def _escape_xml(text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    def _apply_inline_formatting(text: str) -> str:
        escaped: str = _escape_xml(text)
        escaped = _re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", escaped)
        escaped = _re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
        escaped = _re.sub(r"__(.+?)__", r"<b>\1</b>", escaped)
        escaped = _re.sub(r"\*(.+?)\*", r"<i>\1</i>", escaped)
        escaped = _re.sub(r"`(.+?)`", rf"<font face='{BODY_FONT_BOLD}' color='#4004DA'>\1</font>", escaped)
        return escaped

    def _parse_md_table(lines: list) -> tuple:
        headers: list = []
        rows: list = []
        for line in lines:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(_re.match(r"^[-:]+$", c) for c in cells):
                continue
            if not headers:
                headers = cells
            else:
                rows.append(cells)
        return headers, rows

    # ── Document Setup ───────────────────────────────────────────────
    buffer = _io.BytesIO()
    page_w, page_h = A4

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=22 * mm,
        rightMargin=22 * mm,
        topMargin=20 * mm,
        bottomMargin=28 * mm,  # Extra space for page number footer
    )

    # Page number & branded footer drawn on every page
    def _draw_page_footer(canvas, doc_obj):
        canvas.saveState()
        # Thin hairline divider
        canvas.setStrokeColor(NOMU_BORDER)
        canvas.setLineWidth(0.5)
        canvas.line(22 * mm, 18 * mm, page_w - 22 * mm, 18 * mm)
        # Left: brand text
        canvas.setFont(BODY_FONT, 7)
        canvas.setFillColor(NOMU_MUTED)
        canvas.drawString(22 * mm, 13.5 * mm, "Nomu Compliance Flow")
        # Right: page number
        page_num = canvas.getPageNumber()
        canvas.drawRightString(page_w - 22 * mm, 13.5 * mm, f"Page {page_num}")
        canvas.restoreState()

    # ── Styles ───────────────────────────────────────────────────────
    title_style = ParagraphStyle(
        "ReportTitle",
        fontName=HEADING_FONT_BOLD,
        fontSize=24,
        leading=30,
        spaceAfter=4,
        textColor=NOMU_WHITE,
        alignment=TA_LEFT,
    )
    subtitle_style = ParagraphStyle(
        "ReportSubtitle",
        fontName=BODY_FONT,
        fontSize=11,
        leading=15,
        textColor=HexColor("#E0D4F7"),
        alignment=TA_LEFT,
    )
    h1_style = ParagraphStyle(
        "H1",
        fontName=HEADING_FONT_BOLD,
        fontSize=17,
        leading=22,
        spaceBefore=16,
        spaceAfter=6,
        textColor=NOMU_PRIMARY,
    )
    h2_style = ParagraphStyle(
        "H2",
        fontName=HEADING_FONT_BOLD,
        fontSize=14,
        leading=19,
        spaceBefore=14,
        spaceAfter=5,
        textColor=NOMU_BLACK,
    )
    h3_style = ParagraphStyle(
        "H3",
        fontName=HEADING_FONT_BOLD,
        fontSize=11.5,
        leading=16,
        spaceBefore=12,
        spaceAfter=4,
        textColor=NOMU_DARK,
    )
    h4_style = ParagraphStyle(
        "H4",
        fontName=HEADING_FONT,
        fontSize=11,
        leading=16,
        spaceBefore=10,
        spaceAfter=4,
        textColor=NOMU_MUTED,
        fontStyle="italic",
    )
    body_style = ParagraphStyle(
        "ReportBody",
        fontName=BODY_FONT,
        fontSize=9.5,
        leading=15,
        spaceAfter=5,
        textColor=NOMU_BLACK,
    )
    bullet_style = ParagraphStyle(
        "ReportBullet",
        fontName=BODY_FONT,
        fontSize=9.5,
        leading=15,
        leftIndent=14,
        bulletIndent=4,
        spaceAfter=3,
        textColor=NOMU_BLACK,
    )
    numbered_style = ParagraphStyle(
        "ReportNumbered",
        fontName=BODY_FONT,
        fontSize=9.5,
        leading=15,
        leftIndent=14,
        spaceAfter=3,
        textColor=NOMU_BLACK,
    )
    table_cell_style = ParagraphStyle(
        "TableCell",
        fontName=BODY_FONT,
        fontSize=9,
        leading=13,
        textColor=NOMU_BLACK,
    )
    table_header_style = ParagraphStyle(
        "TableHeader",
        fontName=BODY_FONT_BOLD,
        fontSize=9,
        leading=13,
        textColor=NOMU_BLACK,
    )

    # ── Build Story ──────────────────────────────────────────────────
    story: list = []
    title_str: str = title or "Compliance Report"

    # ── Title header block (colored background) ──────────────────────
    import datetime as _dt
    date_str: str = _dt.datetime.now().strftime("%B %d, %Y")

    header_data = [[
        Paragraph(_escape_xml(title_str), title_style),
    ], [
        Paragraph(f"Generated {date_str}", subtitle_style),
    ]]
    available_width: float = page_w - 44 * mm
    header_table = Table(
        header_data,
        colWidths=[available_width],
    )
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), NOMU_PRIMARY),
        ("TOPPADDING", (0, 0), (-1, 0), 18),
        ("BOTTOMPADDING", (0, -1), (-1, -1), 16),
        ("LEFTPADDING", (0, 0), (-1, -1), 22),
        ("RIGHTPADDING", (0, 0), (-1, -1), 22),
        ("TOPPADDING", (0, 1), (-1, 1), 2),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 12 * mm))

    h1_seen: bool = False  # Track whether we've seen the first H1

    # ── Parse content ────────────────────────────────────────────────
    lines = content.split("\n")
    i: int = 0
    while i < len(lines):
        line: str = lines[i]
        stripped: str = line.strip()

        if not stripped:
            story.append(Spacer(1, 3 * mm))
            i += 1
            continue

        # Table detection
        if "|" in stripped and i + 1 < len(lines) and _re.search(r"[-|:]{3,}", lines[i + 1]):
            table_lines: list = []
            while i < len(lines) and "|" in lines[i].strip():
                table_lines.append(lines[i])
                i += 1

            headers, rows = _parse_md_table(table_lines)
            if headers:
                col_count: int = len(headers)
                tbl_data: list = [
                    [Paragraph(_escape_xml(h), table_header_style) for h in headers]
                ]
                for row in rows:
                    padded = row + [""] * (col_count - len(row))
                    tbl_data.append(
                        [Paragraph(_apply_inline_formatting(c), table_cell_style) for c in padded[:col_count]]
                    )

                tbl_width: float = page_w - 44 * mm
                col_widths: list = [tbl_width / col_count] * col_count

                t = Table(tbl_data, colWidths=col_widths, repeatRows=1)
                t.setStyle(TableStyle([
                    # Header row — light background, not heavy
                    ("BACKGROUND", (0, 0), (-1, 0), NOMU_PRIMARY_LIGHT),
                    ("TEXTCOLOR", (0, 0), (-1, 0), NOMU_BLACK),
                    ("FONTSIZE", (0, 0), (-1, 0), 9),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                    ("TOPPADDING", (0, 0), (-1, 0), 8),
                    # Body rows
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [NOMU_WHITE, NOMU_TABLE_ALT]),
                    ("BOTTOMPADDING", (0, 1), (-1, -1), 7),
                    ("TOPPADDING", (0, 1), (-1, -1), 7),
                    # Borders — clean hairlines only
                    ("LINEBELOW", (0, 0), (-1, 0), 0.75, NOMU_PRIMARY),
                    ("LINEBELOW", (0, 1), (-1, -2), 0.25, NOMU_BORDER),
                    ("LINEBELOW", (0, -1), (-1, -1), 0.5, NOMU_BORDER),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ]))
                story.append(Spacer(1, 3 * mm))
                story.append(KeepTogether([t]))
                story.append(Spacer(1, 6 * mm))
            continue

        # Headings — H1 gets a page break, H2 gets an underline accent
        if stripped.startswith("#### "):
            story.append(Paragraph(_apply_inline_formatting(stripped[5:]), h4_style))
        elif stripped.startswith("### "):
            story.append(Paragraph(_apply_inline_formatting(stripped[4:]), h3_style))
        elif stripped.startswith("## "):
            story.append(Paragraph(_apply_inline_formatting(stripped[3:]), h2_style))
            story.append(HRFlowable(
                width="20%", thickness=0.5, color=NOMU_BORDER,
                spaceBefore=1, spaceAfter=5,
            ))
        elif stripped.startswith("# "):
            if h1_seen:
                story.append(PageBreak())
            h1_seen = True
            story.append(Paragraph(_apply_inline_formatting(stripped[2:]), h1_style))
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• "):
            bullet_text: str = stripped[2:] if stripped[0] in "-*" else stripped[2:]
            story.append(Paragraph(
                f"<font color='#9CA3AF' size='7'>•</font>  {_apply_inline_formatting(bullet_text)}",
                bullet_style,
            ))
        # Numbered lists
        elif _re.match(r"^\d+\.\s", stripped):
            num_match = _re.match(r"^(\d+)\.\s(.+)", stripped)
            if num_match:
                num, text = num_match.group(1), num_match.group(2)
                story.append(Paragraph(
                    f"<font color='#374151'><b>{num}.</b></font>  {_apply_inline_formatting(text)}",
                    numbered_style,
                ))
            else:
                story.append(Paragraph(_apply_inline_formatting(stripped), numbered_style))
        # Horizontal rule
        elif _re.match(r"^[-*_]{3,}$", stripped):
            story.append(HRFlowable(
                width="100%", thickness=0.5, color=NOMU_BORDER,
                spaceBefore=4, spaceAfter=4,
            ))
        # Regular paragraph
        else:
            story.append(Paragraph(_apply_inline_formatting(stripped), body_style))

        i += 1

    # ── End-of-report marker ──────────────────────────────────────────
    story.append(Spacer(1, 12 * mm))
    story.append(HRFlowable(
        width="30%", thickness=0.5, color=NOMU_BORDER,
        spaceBefore=0, spaceAfter=6,
    ))
    end_style = ParagraphStyle(
        "EndMark",
        fontName=HEADING_FONT_BOLD,
        fontSize=9,
        textColor=NOMU_MUTED,
        alignment=TA_CENTER,
    )
    story.append(Paragraph("— End of Report —", end_style))

    doc.build(story, onFirstPage=_draw_page_footer, onLaterPages=_draw_page_footer)
    return buffer.getvalue()


def generate_docx(content: str, title: str) -> bytes:
    """Convert markdown content to a styled DOCX using python-docx.

    Pure function. No side effects.

    Args:
        content: Markdown content body.
        title: Report title.

    Returns:
        DOCX file as bytes.
    """
    import io as _io
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    title_str: str = title or "Compliance Report"
    doc.add_heading(title_str, level=0)

    # Parse markdown lines into styled document elements
    for line in content.split("\n"):
        stripped: str = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("```"):
            continue  # Skip code fence markers
        else:
            doc.add_paragraph(stripped)

    buffer = _io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_xlsx(content: str, title: str, input_data: Dict[str, Any]) -> bytes:
    """Create a workbook with Summary and Data sheets.

    Pure function. No side effects.

    Args:
        content: AI-generated summary text for the Summary sheet.
        title: Report title.
        input_data: Raw upstream data for the Data sheet.

    Returns:
        XLSX file as bytes.
    """
    import io as _io
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment

    wb = Workbook()

    # --- Summary sheet ---
    ws_summary = wb.active
    ws_summary.title = "Summary"

    title_str: str = title or "Compliance Report"
    ws_summary["A1"] = title_str
    ws_summary["A1"].font = Font(size=16, bold=True)

    # Write content line by line starting from row 3
    for row_idx, line in enumerate(content.split("\n"), start=3):
        ws_summary.cell(row=row_idx, column=1, value=line)
    ws_summary.column_dimensions["A"].width = 80

    # --- Data sheet ---
    ws_data = wb.create_sheet("Data")
    ws_data["A1"] = "Key"
    ws_data["B1"] = "Value"
    ws_data["A1"].font = Font(bold=True)
    ws_data["B1"].font = Font(bold=True)

    row_idx = 2
    for key, value in input_data.items():
        ws_data.cell(row=row_idx, column=1, value=str(key))
        ws_data.cell(row=row_idx, column=2, value=str(value)[:1000])
        row_idx += 1

    ws_data.column_dimensions["A"].width = 30
    ws_data.column_dimensions["B"].width = 60

    buffer = _io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()
