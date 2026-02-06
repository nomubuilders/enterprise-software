"""Document parsing service for PDF, DOCX, and TXT files."""

import re
from io import BytesIO
from typing import Optional

import PyPDF2
import docx


def parse_pdf(file_content: bytes) -> dict:
    """Extract text from PDF file bytes."""
    try:
        reader = PyPDF2.PdfReader(BytesIO(file_content))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)

        raw_text = "\n".join(pages)
        cleaned_text = clean_text(raw_text)

        return {
            "success": True,
            "text": cleaned_text,
            "metadata": {
                "pages": len(reader.pages),
                "type": "pdf",
            }
        }
    except Exception as e:
        return {"success": False, "text": "", "metadata": {}, "error": str(e)}


def parse_docx(file_content: bytes) -> dict:
    """Extract text from DOCX file bytes."""
    try:
        doc = docx.Document(BytesIO(file_content))
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        raw_text = "\n".join(paragraphs)
        cleaned_text = clean_text(raw_text)

        return {
            "success": True,
            "text": cleaned_text,
            "metadata": {
                "pages": len(doc.paragraphs) // 40 + 1,  # rough estimate
                "type": "docx",
            }
        }
    except Exception as e:
        return {"success": False, "text": "", "metadata": {}, "error": str(e)}


def parse_txt(file_content: bytes, encoding: str = "utf-8") -> dict:
    """Extract text from plain text file bytes."""
    try:
        # Try UTF-8 first, then fall back to latin-1
        try:
            text = file_content.decode(encoding)
        except UnicodeDecodeError:
            text = file_content.decode("latin-1")

        cleaned_text = clean_text(text)
        lines = cleaned_text.split("\n")

        return {
            "success": True,
            "text": cleaned_text,
            "metadata": {
                "pages": len(lines) // 50 + 1,  # rough estimate
                "type": "txt",
            }
        }
    except Exception as e:
        return {"success": False, "text": "", "metadata": {}, "error": str(e)}


def clean_text(text: str) -> str:
    """Clean extracted text: normalize whitespace, remove page numbers."""
    # Remove extra whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    # Remove standalone page numbers (line with just a number)
    text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
    # Normalize line breaks (3+ newlines -> 2)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Strip leading/trailing whitespace
    text = text.strip()
    return text


def parse_document(file_content: bytes, filename: str) -> dict:
    """Parse a document based on file extension."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        result = parse_pdf(file_content)
    elif ext in ("docx", "doc"):
        result = parse_docx(file_content)
    elif ext in ("txt", "text", "md"):
        result = parse_txt(file_content)
    else:
        return {
            "success": False,
            "text": "",
            "metadata": {},
            "error": f"Unsupported file type: .{ext}. Supported: PDF, DOCX, TXT"
        }

    # Add file size to metadata
    if result.get("metadata"):
        result["metadata"]["size"] = len(file_content)
        result["metadata"]["filename"] = filename

    return result
